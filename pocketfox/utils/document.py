"""Document text extraction utilities for LLM context."""

import base64
import csv
import io
import json
import zipfile
from html.parser import HTMLParser
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

from loguru import logger

DOCUMENT_SUFFIXES: set[str] = {
    ".pdf", ".docx", ".csv", ".txt", ".html", ".htm",
    ".odt", ".rtf", ".epub", ".json", ".xlsx",
}

IMAGE_SUFFIXES: set[str] = {".jpg", ".jpeg", ".png", ".gif", ".webp"}

# Hard cap on extracted text to avoid blowing up LLM context
_MAX_TEXT_CHARS = 200_000


def extract_document_text(path: Path, max_bytes: int = 10 * 1024 * 1024) -> str | None:
    """Extract text content from a document file.

    Args:
        path: Path to the document.
        max_bytes: Maximum file size in bytes.

    Returns:
        Extracted text, or None if unsupported/too large/error.
    """
    if not path.is_file():
        return None

    size = path.stat().st_size
    if size > max_bytes:
        logger.warning(
            f"Document {path.name} too large "
            f"({size / 1024 / 1024:.1f} MiB > {max_bytes / 1024 / 1024:.1f} MiB)"
        )
        return None

    ext = path.suffix.lower()
    handler = _EXTRACTORS.get(ext)
    if handler is None:
        return None

    try:
        text = handler(path)
    except Exception as exc:
        logger.warning(f"Failed to extract text from {path.name}: {exc}")
        return None

    if not text or not text.strip():
        logger.info(f"Document {path.name} produced empty text")
        return None

    text = text.strip()
    if len(text) > _MAX_TEXT_CHARS:
        text = text[:_MAX_TEXT_CHARS] + f"\n\n[truncated at {_MAX_TEXT_CHARS:,} characters]"

    return text


_DOCUMENT_MIME: dict[str, str] = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".csv": "text/csv",
    ".txt": "text/plain",
    ".html": "text/html",
    ".htm": "text/html",
    ".json": "application/json",
    ".odt": "application/vnd.oasis.opendocument.text",
    ".rtf": "application/rtf",
    ".epub": "application/epub+zip",
}


def encode_document_block(path: Path, max_bytes: int = 10 * 1024 * 1024) -> dict[str, Any] | None:
    """Encode a document as a MIME-annotated base64 content block.

    Produces a ``document`` content block compatible with Anthropic's API.
    No external library needed — just base64 encoding of the raw file.

    Args:
        path: Path to the document file.
        max_bytes: Maximum file size in bytes.

    Returns:
        Content block dict, or None if too large or unsupported extension.
    """
    if not path.is_file():
        return None

    mime = _DOCUMENT_MIME.get(path.suffix.lower())
    if not mime:
        return None

    size = path.stat().st_size
    if size > max_bytes:
        return None

    raw = path.read_bytes()
    b64 = base64.b64encode(raw).decode("ascii")
    return {
        "type": "document",
        "source": {
            "type": "base64",
            "media_type": mime,
            "data": b64,
        },
    }


# Keep backward-compat alias
encode_pdf_block = encode_document_block


# ---------------------------------------------------------------------------
# Per-format extractors
# ---------------------------------------------------------------------------


def _extract_txt(path: Path) -> str | None:
    """Plain text — try UTF-8, fall back to latin-1."""
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1")


def _extract_csv(path: Path) -> str | None:
    """CSV — read as text (LLMs can parse tabular CSV directly)."""
    try:
        raw = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        raw = path.read_text(encoding="latin-1")
    # Validate it parses as CSV (catches garbage files)
    reader = csv.reader(io.StringIO(raw))
    rows = list(reader)
    if not rows:
        return None
    return raw


def _extract_json(path: Path) -> str | None:
    """JSON — pretty-print for readability."""
    raw = path.read_text(encoding="utf-8")
    data = json.loads(raw)
    return json.dumps(data, indent=2, ensure_ascii=False)


def _extract_html(path: Path) -> str | None:
    """HTML — extract readable text via readability-lxml (already a dependency)."""
    try:
        raw = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        raw = path.read_text(encoding="latin-1")
    try:
        from lxml.html.clean import Cleaner
        from readability import Document

        doc = Document(raw)
        summary = doc.summary()
        cleaner = Cleaner(scripts=True, javascript=True, style=True)
        cleaned = cleaner.clean_html(summary)

        # Strip remaining tags
        text = _strip_html_tags(cleaned)
        return text if text.strip() else _strip_html_tags(raw)
    except Exception:
        # Fallback: simple tag stripping
        return _strip_html_tags(raw)


def _extract_pdf(path: Path) -> str | None:
    """PDF — extract text via pymupdf (optional dependency)."""
    try:
        import pymupdf  # noqa: F811
    except ImportError:
        logger.warning(
            "pymupdf not installed; PDF text extraction unavailable. "
            "Install with: pip install pymupdf"
        )
        return None

    doc = pymupdf.open(str(path))
    pages = []
    for page in doc:
        text = page.get_text()
        if text.strip():
            pages.append(text)
    doc.close()
    return "\n\n".join(pages) if pages else None


def _extract_docx(path: Path) -> str | None:
    """DOCX — extract paragraph text via python-docx (optional dependency)."""
    try:
        from docx import Document
    except ImportError:
        logger.warning(
            "python-docx not installed; .docx files not supported. "
            "Install with: pip install python-docx"
        )
        return None

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs) if paragraphs else None


def _extract_xlsx(path: Path) -> str | None:
    """XLSX — extract sheet data as CSV-like text via openpyxl (optional dependency)."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        logger.warning(
            "openpyxl not installed; .xlsx files not supported. "
            "Install with: pip install openpyxl"
        )
        return None

    wb = load_workbook(str(path), read_only=True, data_only=True)
    sections = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                rows.append(",".join(cells))
        if rows:
            header = f"[Sheet: {sheet_name}]" if len(wb.sheetnames) > 1 else ""
            sections.append(f"{header}\n{chr(10).join(rows)}" if header else "\n".join(rows))
    wb.close()
    return "\n\n".join(sections) if sections else None


def _extract_odt(path: Path) -> str | None:
    """ODT — extract text from OpenDocument Text (stdlib zipfile + xml)."""
    with zipfile.ZipFile(path, "r") as zf:
        if "content.xml" not in zf.namelist():
            return None
        content = zf.read("content.xml")

    ns = {"text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0"}
    tree = ElementTree.fromstring(content)
    paragraphs = []
    for p in tree.iter(f"{{{ns['text']}}}p"):
        text = "".join(p.itertext())
        if text.strip():
            paragraphs.append(text)
    return "\n\n".join(paragraphs) if paragraphs else None


def _extract_rtf(path: Path) -> str | None:
    """RTF — extract text via striprtf (optional dependency)."""
    try:
        from striprtf.striprtf import rtf_to_text
    except ImportError:
        logger.warning(
            "striprtf not installed; .rtf files not supported. "
            "Install with: pip install striprtf"
        )
        return None

    raw = path.read_text(encoding="latin-1")
    return rtf_to_text(raw)


def _extract_epub(path: Path) -> str | None:
    """EPUB — extract text from XHTML chapters (stdlib zipfile + html.parser)."""
    with zipfile.ZipFile(path, "r") as zf:
        parts = []
        for name in sorted(zf.namelist()):
            if name.endswith((".xhtml", ".html", ".htm")):
                raw = zf.read(name).decode("utf-8", errors="replace")
                text = _strip_html_tags(raw)
                if text.strip():
                    parts.append(text.strip())
    return "\n\n".join(parts) if parts else None


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


class _TagStripper(HTMLParser):
    """Minimal HTML tag stripper."""

    def __init__(self):
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        self.parts.append(data)


def _strip_html_tags(html: str) -> str:
    stripper = _TagStripper()
    stripper.feed(html)
    return "".join(stripper.parts)


# Dispatch table: extension → extractor function
_EXTRACTORS: dict[str, Any] = {
    ".txt": _extract_txt,
    ".csv": _extract_csv,
    ".json": _extract_json,
    ".html": _extract_html,
    ".htm": _extract_html,
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".xlsx": _extract_xlsx,
    ".odt": _extract_odt,
    ".rtf": _extract_rtf,
    ".epub": _extract_epub,
}
